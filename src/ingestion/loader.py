# src/ingestion/loader.py
from typing import List, Dict, Any
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import pandas as pd
from loguru import logger


class DocumentLoader:
    """Load documents from various file formats"""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".csv", ".xlsx"}

    def load_file(self, file_path: str) -> Dict[str, Any]:
        """Load a single file"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {path.suffix}")

        logger.info(f"Loading file: {path.name}")

        if path.suffix == ".pdf":
            return self._load_pdf(path)
        elif path.suffix == ".txt" or path.suffix == ".md":
            return self._load_text(path)
        elif path.suffix == ".docx":
            return self._load_docx(path)
        elif path.suffix == ".csv":
            return self._load_csv(path)
        elif path.suffix == ".xlsx":
            return self._load_excel(path)

    def _load_pdf(self, path: Path) -> Dict[str, Any]:
        """Load PDF file"""
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"

        return {
            "content": text,
            "metadata": {
                "source": str(path),
                "filename": path.name,
                "type": "pdf",
                "pages": len(reader.pages),
            },
        }

    def _load_text(self, path: Path) -> Dict[str, Any]:
        """Load text file"""
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        return {
            "content": text,
            "metadata": {"source": str(path), "filename": path.name, "type": "text"},
        }

    def _load_docx(self, path: Path) -> Dict[str, Any]:
        """Load DOCX file"""
        doc = Document(path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        return {
            "content": text,
            "metadata": {
                "source": str(path),
                "filename": path.name,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
            },
        }

    def _load_csv(self, path: Path) -> Dict[str, Any]:
        """Load CSV file"""
        df = pd.read_csv(path)
        text = df.to_string()

        return {
            "content": text,
            "metadata": {
                "source": str(path),
                "filename": path.name,
                "type": "csv",
                "rows": len(df),
                "columns": list(df.columns),
            },
        }

    def _load_excel(self, path: Path) -> Dict[str, Any]:
        """Load Excel file"""
        df = pd.read_excel(path)
        text = df.to_string()

        return {
            "content": text,
            "metadata": {
                "source": str(path),
                "filename": path.name,
                "type": "xlsx",
                "rows": len(df),
                "columns": list(df.columns),
            },
        }

    def load_directory(self, directory: str) -> List[Dict[str, Any]]:
        """Load all supported files from directory"""
        documents = []
        dir_path = Path(directory)

        for file_path in dir_path.rglob("*"):
            if file_path.is_file() and file_path.suffix in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
